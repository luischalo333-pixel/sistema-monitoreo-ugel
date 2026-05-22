import streamlit as st
from google import genai
from google.genai import types
from PIL import Image
from datetime import date, time
import io
import pandas as pd
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fpdf import FPDF

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="SGA-P: UGEL LA UNIÓN", layout="wide", page_icon="📋")

# --- ESTILOS VISUALES DE LA APP ---
st.markdown("""
    <style>
    .main { background-color: #f8fafc; }
    .title-container { 
        background-color: #1e3a8a; padding: 25px; border-radius: 12px; 
        color: white; text-align: center; border-bottom: 5px solid #facc15;
        box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1);
    }
    .section-header {
        background-color: #e2e8f0; padding: 10px 14px; border-radius: 6px;
        color: #1e3a8a; font-weight: bold; margin-top: 15px; margin-bottom: 10px;
        border-left: 5px solid #1e3a8a;
    }
    </style>
    """, unsafe_allow_html=True)

# --- BARRA LATERAL ---
with st.sidebar:
    st.header("🔐 SEGURIDAD")
    api_key_input = st.text_input("Ingrese su Google Gemini API Key:", type="password")
    if api_key_input:
        st.success("✅ Clave ingresada")
    else:
        st.error("⚠️ Se requiere ingresar una clave API para usar la IA.")

# --- ENCABEZADO ---
st.markdown("""
    <div class="title-container">
        <h1 style='margin:0; font-size:26px; font-weight:bold;'>SISTEMA DE GESTIÓN DEL ACOMPAÑAMIENTO PEDAGÓGICO</h1>
        <h3 style='margin:5px; font-size:16px; color:#facc15;'>MONITOREO DE PRÁCTICAS DOCENTES - UGEL LA UNIÓN</h3>
    </div>
    """, unsafe_allow_html=True)

# --- INICIALIZAR SESSION STATE ---
if 'analisis' not in st.session_state: st.session_state.analisis = ""
if 'datos' not in st.session_state: st.session_state.datos = {}
if 'compromisos' not in st.session_state: st.session_state.compromisos = ""
if 'analisis_realizado' not in st.session_state: st.session_state.analisis_realizado = False

t1, t2, t3, t4 = st.tabs(["📋 REGISTRO", "📊 ANÁLISIS IA", "💡 ASESORÍA", "📥 DESCARGAS"])

# ==========================================
# PESTAÑA 1: REGISTRO
# ==========================================
with t1:
    st.subheader("🖊️ Registro de Datos Generales del Monitoreo")
    with st.form("form_ugel_detallado"):
        st.markdown("<div class='section-header'>🏢 1. Datos Generales de la I.E.</div>", unsafe_allow_html=True)
        g1, g2, g3 = st.columns(3)
        with g1:
            ie_ugel = st.text_input("UGEL", value="La Unión")
            ie_nombre = st.text_input("Número / Nombre de la I.E.")
        with g2:
            ie_modalidad = st.selectbox("Modalidad", ["EBR (Regular)", "EIB (Intercultural Bilingüe)", "EBA (Alternativa)", "EBE (Especial)"])
            ie_niveles = st.text_input("Niveles que atiende la I.E.")
        with g3:
            ie_cod_modular = st.text_input("Código Modular")
            ie_fecha_mon = st.date_input("Fecha de Monitoreo", date.today())

        st.markdown("<div class='section-header'>👨‍💼 2. Datos del Director</div>", unsafe_allow_html=True)
        d1, d2 = st.columns(2)
        with d1:
            dir_nom = st.text_input("Nombres y Apellidos (Director)")
            dir_cel = st.text_input("N° de Celular (Director)")
        with d2:
            dir_dni = st.text_input("DNI (Director)")
            dir_correo = st.text_input("Correo Electrónico (Director)")

        st.markdown("<div class='section-header'>👩‍🏫 3. Datos del Docente Observado</div>", unsafe_allow_html=True)
        doc1, doc2, doc3 = st.columns(3)
        with doc1:
            doc_nom = st.text_input("Nombres y Apellidos (Docente)")
            doc_dni = st.text_input("DNI (Docente)")
            doc_area = st.text_input("Área Curricular")
        with doc2:
            doc_grado = st.text_input("Grado y Sección")
            doc_correo = st.text_input("Correo Electrónico (Docente)")
            doc_cel = st.text_input("N° de Celular (Docente)")
        with doc3:
            doc_nivel = st.selectbox("Nivel Educativo", ["Inicial", "Primaria", "Secundaria"])
            doc_esp = st.text_input("Especialidad")

        st.markdown("<div class='section-header'>🧐 4. Datos del Especialista Monitor</div>", unsafe_allow_html=True)
        esp1, esp2, esp3 = st.columns(3)
        with esp1:
            esp_nom = st.text_input("Nombres y Apellidos (Especialista)")
            esp_dni = st.text_input("DNI (Especialista)")
        with esp2:
            esp_cel = st.text_input("N° de Celular (Especialista)")
            esp_correo = st.text_input("Correo Electrónico (Especialista)")
        with esp3:
            esp_visita = st.selectbox("Número de Visita", ["1ra Visita", "2da Visita", "3ra Visita", "4ta Visita"])
            esp_h_inicio = st.time_input("Hora de Inicio", time(8, 0))
            esp_h_term = st.time_input("Hora de Término", time(10, 0))
            esp_fecha_app = st.date_input("Fecha de Aplicación", date.today())

        if st.form_submit_button("💾 GUARDAR TODOS LOS DATOS"):
            st.session_state.datos = {
                "UGEL": ie_ugel, "IE_Nombre": ie_nombre, "IE_Modalidad": ie_modalidad, "IE_Niveles": ie_niveles, "IE_Codigo_Modular": ie_cod_modular, "IE_Fecha_Monitoreo": str(ie_fecha_mon),
                "Director_Nombre": dir_nom, "Director_DNI": dir_dni, "Director_Celular": dir_cel, "Director_Correo": dir_correo,
                "Docente_Nombre": doc_nom, "Docente_DNI": doc_dni, "Docente_Area": doc_area, "Docente_Grado_Seccion": doc_grado, "Docente_Correo": doc_correo, "Docente_Celular": doc_cel, "Docente_Nivel": doc_nivel, "Docente_Especialidad": doc_esp,
                "Especialista_Nombre": esp_nom, "Especialista_DNI": esp_dni, "Especialista_Celular": esp_cel, "Especialista_Correo": esp_correo, "Especialista_Visita": esp_visita, "Especialista_Hora_Inicio": str(esp_h_inicio), "Especialista_Hora_Termino": str(esp_h_term), "Especialista_Fecha_Aplicacion": str(esp_fecha_app)
            }
            st.success("✅ Datos guardados con éxito. Ya puedes realizar el análisis.")

# ==========================================
# PESTAÑA 2: ANÁLISIS IA Y DESCARGAS
# ==========================================
with t2:
    st.subheader("📊 Análisis Avanzado de Evidencias con Gemini")
    archivos = st.file_uploader("Subir evidencias (PDF/Word/Imágenes)", accept_multiple_files=True, type=["pdf", "docx", "txt", "jpg", "jpeg", "png"], key="uploader_final")
    
    if st.button("🔍 INICIAR ANÁLISIS"):
        if not api_key_input:
            st.error("❌ Por favor, ingresa tu API Key en la barra lateral.")
        elif not archivos:
            st.warning("⚠️ Sube al menos un archivo para el análisis.")
        else:
            with st.spinner("🚀 Analizando pedagógicamente..."):
                try:
                    client = genai.Client(api_key=api_key_input)
                    prompt_base = """Actúa como un Especialista en Acompañamiento Pedagógico de UGEL. 
                    Analiza de forma exhaustiva los siguientes documentos y evidencias del docente adjuntos. 
                    Evalúa el desempeño basándote rigurosamente en las Rúbricas de Observación de Aula del Marco de Buen Desempeño Docente del Minedu.
                    Estructura tu respuesta estrictamente con:
                    1. Logros identificados (Puntos fuertes).
                    2. Necesidades de formación o aspectos prioritarios a mejorar.
                    3. Sugerencias detalladas de práctica pedagógica orientadas a la mejora continua."""
                    
                    contenidos = [prompt_base]
                    for f in archivos:
                        if f.type == "application/pdf":
                            pdf_reader = PdfReader(f)
                            texto_pdf = "".join([pagina.extract_text() or "" for pagina in pdf_reader.pages])
                            contenidos.append(f"--- TEXTO PDF ({f.name}) ---\n{texto_pdf}")
                        elif f.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or f.name.endswith(".docx"):
                            doc_memoria = Document(io.BytesIO(f.read()))
                            lista_texto = [p.text for p in doc_memoria.paragraphs if p.text.strip()]
                            contenidos.append(f"--- TEXTO WORD ({f.name}) ---\n" + "\n".join(lista_texto))
                        elif f.type in ["image/jpeg", "image/png"]:
                            contenidos.append(Image.open(io.BytesIO(f.read())))
                        elif f.type == "text/plain":
                            contenidos.append(f.read().decode('utf-8', errors='ignore'))
                    
                    try:
                        respuesta = client.models.generate_content(model='gemini-2.5-flash', contents=contenidos)
                    except Exception as inner_error:
                        if "503" in str(inner_error) or "UNAVAILABLE" in str(inner_error):
                            st.warning("⏳ Canal ocupado, redirigiendo a alta prioridad...")
                            respuesta = client.models.generate_content(model='gemini-2.5-pro', contents=contenidos)
                        else:
                            raise inner_error
                    
                    st.session_state.analisis = respuesta.text
                    st.session_state.analisis_realizado = True
                    st.success("✨ ¡Análisis completado!")
                except Exception as e:
                    st.error(f"❌ Error: {e}")

    if st.session_state.analisis_realizado:
        st.markdown("### 📝 Resultado del Análisis")
        st.write(st.session_state.analisis)
        
        st.markdown("---")
        st.subheader("📥 EXPORTAR INFORME OFICIAL")
        st.info("Descarga el documento con el formato oficial institucional con un solo clic:")
        
        c1, c2 = st.columns(2)
        doc_nom_raw = st.session_state.datos.get('Docente_Nombre', 'Docente')
        nombre_archivo = doc_nom_raw.replace(" ", "_") if doc_nom_raw else "Informe"

        # --- GENERADOR WORD FORMATO PREMIUM ---
        with c1:
            try:
                doc = Document()
                # Márgenes profesionales
                for section in doc.sections:
                    section.top_margin = Inches(1)
                    section.bottom_margin = Inches(1)
                    section.left_margin = Inches(1)
                    section.right_margin = Inches(1)

                # Encabezado principal
                p_t = doc.add_paragraph()
                p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_t = p_t.add_run("INFORME DE ACOMPAÑAMIENTO Y MONITOREO PEDAGÓGICO\n")
                r_t.bold = True
                r_t.font.size = Pt(14)
                r_t.font.name = 'Arial'
                r_t.font.color.rgb = RGBColor(30, 58, 138) # Azul Institucional
                
                r_sub = p_t.add_run("UGEL LA UNIÓN - DIRECCIÓN REGIONAL DE EDUCACIÓN AREQUIPA")
                r_sub.font.size = Pt(10)
                r_sub.font.name = 'Arial'
                r_sub.font.color.rgb = RGBColor(100, 116, 139)

                # Tabla de Datos Generales si existen
                if st.session_state.datos:
                    doc.add_heading("I. DATOS INFORMATIVOS", level=1)
                    table = doc.add_table(rows=1, cols=2)
                    table.style = 'Light Shading Accent 1'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Campo Técnico'
                    hdr_cells[1].text = 'Información Registrada'
                    
                    datos_tabla = [
                        ("Institución Educativa", st.session_state.datos.get('IE_Nombre')),
                        ("Modalidad / Nivel", f"{st.session_state.datos.get('IE_Modalidad')} - {st.session_state.datos.get('Docente_Nivel')}"),
                        ("Docente Observado", st.session_state.datos.get('Docente_Nombre')),
                        ("Área / Especialidad", f"{st.session_state.datos.get('Docente_Area')} / {st.session_state.datos.get('Docente_Especialidad')}"),
                        ("Especialista Monitor", st.session_state.datos.get('Especialista_Nombre')),
                        ("Fecha de Aplicación", st.session_state.datos.get('Especialista_Fecha_Aplicacion')),
                    ]
                    for campo, valor in datos_tabla:
                        row_cells = table.add_row().cells
                        row_cells[0].text = campo
                        row_cells[1].text = str(valor)

                # Cuerpo del informe
                h2 = doc.add_heading("II. ANÁLISIS DE EVIDENCIAS SEGÚN RÚBRICAS MINEDU", level=1)
                h2.runs[0].font.color.rgb = RGBColor(30, 58, 138)
                
                p_cuerpo = doc.add_paragraph(st.session_state.analisis)
                p_cuerpo.style.font.name = 'Arial'
                p_cuerpo.style.font.size = Pt(11)

                if st.session_state.compromisos:
                    h3 = doc.add_heading("III. ACUERDOS Y COMPROMISOS", level=1)
                    h3.runs[0].font.color.rgb = RGBColor(30, 58, 138)
                    p_comp = doc.add_paragraph(st.session_state.compromisos)
                    p_comp.style.font.name = 'Arial'
                    p_comp.style.font.size = Pt(11)

                out_w = io.BytesIO()
                doc.save(out_w)
                out_w.seek(0)
                
                st.download_button(
                    label="🔷 DESCARGAR EN WORD (.DOCX)",
                    data=out_w.read(),
                    file_name=f"Informe_Monitoreo_{nombre_archivo}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e: st.error(f"Error Word: {e}")

        # --- GENERADOR PDF SEGURO CON FORMATO ELEGANTE ---
        with c2:
            try:
                class PDF(FPDF):
                    def header(self):
                        self.set_fill_color(30, 58, 138) # Fondo azul
                        self.rect(0, 0, 210, 35, 'F')
                        self.set_text_color(255, 255, 255)
                        self.set_font("Helvetica", "B", 14)
                        self.cell(0, 5, "INFORME DE ACOMPAÑAMIENTO PEDAGÓGICO", ln=True, align="C")
                        self.set_font("Helvetica", "", 10)
                        self.cell(0, 5, "UGEL LA UNIÓN - AREQUIPA", ln=True, align="C")
                        self.ln(12)
                    def footer(self):
                        self.set_y(-15)
                        self.set_font("Helvetica", "I", 8)
                        self.set_text_color(150, 150, 150)
                        self.cell(0, 10, f"Página {self.page_no()}", align="C")

                pdf = PDF()
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=20)
                pdf.set_text_color(30, 41, 59) # Gris oscuro para lectura cómoda

                # Sección Informativa en PDF
                if st.session_state.datos:
                    pdf.set_font("Helvetica", "B", 12)
                    pdf.set_text_color(30, 58, 138)
                    pdf.cell(0, 8, "I. DATOS INFORMATIVOS", ln=True)
                    pdf.ln(2)
                    pdf.set_font("Helvetica", "", 10)
                    pdf.set_text_color(50, 50, 50)
                    
                    datos_pdf = [
                        f"Institucion Educativa: {st.session_state.datos.get('IE_Nombre')}",
                        f"Docente Observado: {st.session_state.datos.get('Docente_Nombre')}",
                        f"Area Curricular: {st.session_state.datos.get('Docente_Area')}",
                        f"Especialista Monitor: {st.session_state.datos.get('Especialista_Nombre')}",
                        f"Fecha de Aplicacion: {st.session_state.datos.get('Especialista_Fecha_Aplicacion')}"
                    ]
                    for d in datos_pdf:
                        # Reemplazo seguro para compatibilidad de caracteres planos
                        linea = d.replace("í","i").replace("ó","o").replace("á","a").replace("é","e").replace("ú","u").replace("Ñ","N").replace("ñ","n")
                        pdf.cell(0, 6, f" - {linea}", ln=True)
                    pdf.ln(5)

                # Cuerpo del reporte IA
                pdf.set_font("Helvetica", "B", 12)
                pdf.set_text_color(30, 58, 138)
                pdf.cell(0, 8, "II. ANÁLISIS DE EVIDENCIAS PEDAGÓGICAS", ln=True)
                pdf.ln(2)
                
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(40, 40, 40)
                
                # Tratamiento de strings para evitar rupturas de codificación latin-1
                texto_analisis = st.session_state.analisis
                remplazos = {"á":"a","é":"e","í":"i","ó":"o","ú":"u","ñ":"n","Á":"A","É":"E","Í":"I","Ó":"O","Ú":"U","Ñ":"N"}
                for orig, rep in remplazos.items():
                    texto_analisis = texto_analisis.replace(orig, rep)
                
                pdf.multi_cell(0, 6, texto_analisis)

                out_p = io.BytesIO()
                pdf.output(out_p)
                out_p.seek(0)

                st.download_button(
                    label="🟥 DESCARGAR EN PDF (.PDF)",
                    data=out_p.read(),
                    file_name=f"Informe_Monitoreo_{nombre_archivo}.pdf",
                    mime="application/pdf"
                )
            except Exception as e: st.error(f"Error PDF: {e}")

# ==========================================
# PESTAÑAS ADICIONALES
# ==========================================
with t3:
    st.subheader("💡 Asesoría y Retroalimentación")
    st.session_state.compromisos = st.text_area("Acuerdos y Compromisos asumidos por el Docente:", value=st.session_state.compromisos, height=250)

with t4:
    st.subheader("📥 Descargas Secundarias")
    st.info("Utiliza los paneles superiores de la pestaña 'ANÁLISIS IA' para descargar tus informes profesionales formateados.")